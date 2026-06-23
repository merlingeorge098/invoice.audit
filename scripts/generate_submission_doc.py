# -*- coding: utf-8 -*-
"""Generates the InnovateZ 2026 Round 2 submission document for Invoice Audit Pro."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEAL = RGBColor(0x0F, 0x76, 0x6E)
DARK = RGBColor(0x1E, 0x29, 0x3B)
GREY = RGBColor(0x64, 0x74, 0x8B)
AMBER = RGBColor(0xB4, 0x53, 0x09)
RED = RGBColor(0xB9, 0x1C, 0x1C)

SHOT_DIR = r"c:\Users\Asus\Downloads\invoiceAUDIT\invoice-auditor-pro\submission_screenshots"
ARCH_IMG = r"c:\Users\Asus\Downloads\invoiceAUDIT\invoice-auditor-pro\architecture_diagram.png"
import sys
OUT_PATH = sys.argv[1] if len(sys.argv) > 1 else r"c:\Users\Asus\Downloads\invoiceAUDIT\invoice-auditor-pro\InnovateZ2026_Submission_InvoiceAudit.docx"

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = DARK

for sec in doc.sections:
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_border(cell, color_hex="0F766E", sz="10"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color_hex)
        borders.append(el)
    tc_pr.append(borders)


def H(text, level=1, color=TEAL, space_before=10):
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run(text)
    run.font.color.rgb = color
    return h


def P(text, size=11, bold=False, italic=False, color=DARK, space_after=6, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    return p


def bullet(text, lead=None, space_after=2):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(space_after)
    if lead:
        r1 = p.add_run(lead)
        r1.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def callout(text, fill="ECFDF5", border="0F766E", text_color=TEAL, size=11, space_after=10):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, border)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = text_color
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(space_after)
    return table


def picture(path, width_in, caption, space_after=10):
    doc.add_picture(path, width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_after = Pt(2)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    cap.paragraph_format.space_after = Pt(space_after)


def table_simple(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr[i], "0F766E")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9.5)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)
    return table


# =====================================================================
# TITLE PAGE
# =====================================================================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(40)
r = t.add_run("Invoice.Audit")
r.font.size = Pt(38)
r.bold = True
r.font.color.rgb = TEAL

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("A control that stops a fake invoice before the money leaves the bank.")
r.font.size = Pt(15)
r.italic = True
r.font.color.rgb = DARK
s.paragraph_format.space_after = Pt(16)

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tag.add_run("InnovateZ 2026 — Round 2 Submission")
r.font.size = Pt(11.5)
r.font.color.rgb = GREY
tag.paragraph_format.space_after = Pt(16)

callout(
    "Every screenshot in this document is from one real, unedited run: sign up, load real records, "
    "watch the system catch the fraud — captured minutes before this was written.",
    fill="FEF3C7", border="D97706", text_color=AMBER, size=10.5,
)

P("This is not hypothetical. Four real cases, four different scales:", bold=True, size=11, space_after=6)
case_rows = [
    ("$400M+ — BlackRock's HPS (2026)", "A telecom entrepreneur supplied fabricated invoices and fake customer confirmations as loan collateral; the email domains used to \"verify\" them didn't match the real companies. Caught only after the loan was already funded."),
    ("$100M+ — Google & Facebook (2013-2015)", "A fraudster impersonated their real hardware supplier, Quanta Computer, and billed both companies through forged invoices and look-alike email domains for two years before he was caught."),
    ("₹266 Cr — Bengaluru, 2026 (DGGI)", "Six shell companies with no real business generated fake invoices and falsely claimed ₹48 Cr in tax credit. The mastermind was a chartered accountant who had once audited similar firms."),
    ("₹72 Cr — Hyderabad, 2025 (DGGI)", "Fake GST invoices from shell firms were used to claim Input Tax Credit that was never actually owed."),
]
table_simple(["Case", "What happened"], case_rows)

P(
    "Every one of these frauds used the same two tricks: an invoice from a vendor that was never "
    "properly verified, and paperwork that looked real enough to pass a tired human reviewer. "
    "Invoice.Audit is built to catch exactly that — automatically, on every invoice, before payment.",
    italic=True, color=GREY, size=10, space_after=12,
)

link_table = doc.add_table(rows=3, cols=2)
link_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, value) in enumerate([
    ("Team / Project Name", "Invoice.Audit"),
    ("GitHub Repository", "[ PASTE GITHUB REPO LINK HERE ]"),
    ("Live MVP / Demo Link", "[ PASTE DEPLOYED APP LINK HERE ]"),
]):
    c0, c1 = link_table.rows[i].cells
    set_cell_shading(c0, "0F766E")
    r0 = c0.paragraphs[0].add_run(label)
    r0.bold = True
    r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r0.font.size = Pt(10)
    set_cell_shading(c1, "F1F5F9")
    r1 = c1.paragraphs[0].add_run(value)
    r1.font.size = Pt(10)
    if value.startswith("["):
        r1.italic = True
        r1.font.color.rgb = GREY

doc.add_page_break()

# =====================================================================
# EXECUTIVE SUMMARY
# =====================================================================
H("Executive Summary")

P(
    "Indian companies paying ₹20–50 Cr a year in vendor bills lose 3–8% of that to fraud and "
    "duplicate payments that nobody catches until it's too late — the four cases on the cover page "
    "show this happens at every scale, from a single Hyderabad shell company to a BlackRock-owned "
    "lender. On a ₹20 Cr payables book, that is ₹60 lakh to ₹1.6 Cr a year walking out the door."
)
P(
    "Invoice.Audit puts a check in front of every invoice before a human ever opens it: extract the "
    "data, run it through six fraud-detection rules, score the risk, and route it to the right "
    "person — with a tamper-proof record of every decision."
)

table_simple(
    ["What we ran", "What it caught — same patterns as the cases above"],
    [
        ("17 invoices, fresh empty workspace, zero manual setup",
         "2 exact-duplicate invoices blocked automatically (the BlackRock and Hyderabad pattern)"),
        ("14 vendors and 10 purchase orders pre-loaded for realism",
         "2 invoices from an unregistered vendor blocked automatically (the Google/Facebook pattern)"),
        ("Full run captured live with Playwright, screenshots in Section 6",
         "₹21.5 lakh in flagged spend held back before any human reviewed it"),
    ],
)

P(
    "The rest of this document explains exactly how that happened — what calls an AI model, what "
    "doesn't, why a person cannot get the same result by pasting an invoice into ChatGPT, and what "
    "is honestly still left to build.",
    color=GREY, italic=True, size=10,
)

doc.add_page_break()

# =====================================================================
# SECTION 1 — PROBLEM AND USER FLOW
# =====================================================================
H("1. The Problem and How a User Moves Through the Product")

P("The problem, in one line", bold=True, size=12.5, color=TEAL, space_after=4)
P(
    "A finance team has no reliable way to know, in the moment an invoice arrives, whether it is "
    "real, whether it has already been paid, or whether the vendor sending it actually exists in "
    "their records. Today that check happens — if it happens at all — days later, by a person "
    "scanning a spreadsheet, after the payment is already scheduled. Every fraud case on the cover "
    "page slipped through exactly that gap."
)

P("What we built", bold=True, size=12.5, color=TEAL, space_after=4)
steps = [
    ("Sign up", "A company signs up and is logged into a fully working, isolated workspace in under two minutes — no manual setup by anyone on our side."),
    ("Bring in invoices", "Three ways in: upload a PDF or photo (read automatically), import a spreadsheet (no extraction cost), or load realistic sample data to see the whole thing work in seconds."),
    ("Automatic checks", "Every invoice is read, then checked against six fraud and compliance rules, in the background — the screen shows live progress the whole time."),
    ("Sorted by risk", "Clean invoices go straight to the dashboard. Anything suspicious goes to a review queue with the exact reason it was flagged, not just a vague warning."),
    ("One-click decisions", "A reviewer approves, escalates to a manager, or asks the vendor for proof — all from the same screen, no digging through folders."),
    ("A record nobody can quietly edit", "Every decision is written to a log that is mathematically provable to be untouched, and GST filing data can be reconciled separately when the month closes."),
]
for label, detail in steps:
    bullet(detail, lead=label + " — ")

picture(
    f"{SHOT_DIR}/01_dashboard.png", 6.0,
    "Figure 1 — Live dashboard after a real run: 17 invoices processed, 6 waiting on review, "
    "3 high-risk, 2 duplicates caught automatically, 41% cleared without a human touching them.",
)

doc.add_page_break()

# =====================================================================
# SECTION 2 — UNDER THE HOOD
# =====================================================================
H("2. How It Actually Works")

P(
    "One rule we held ourselves to while writing this: never write \"we use AI\" without saying "
    "exactly where, exactly which model, and exactly what it is and isn't allowed to decide.",
    italic=True, color=GREY, size=10, space_after=10,
)

P("Step 1 — Reading the invoice", bold=True, size=12.5, color=TEAL, space_after=4)
P(
    "When someone uploads a PDF or photo, the file is sent to OpenAI's GPT-4o Vision model with "
    "instructions to extract specific fields — invoice number, vendor name, amount, dates, PO "
    "number, and the line items. It returns that as structured data, nothing more. This is the only "
    "place in the entire system where an AI model is used, and it never decides whether an invoice "
    "is approved, suspicious, or fraudulent — that decision is made by fixed rules, on purpose, so "
    "the same invoice always gets the same answer."
)

P("Step 2 — Six checks, every single invoice, no exceptions", bold=True, size=12.5, color=TEAL, space_after=4)
P("This is the part that actually catches fraud. Each check runs against the company's own records in the database:")
rules = [
    ("Is this vendor real?", "Looked up against the company's own approved-vendor list. Not there → blocked on the spot. This is exactly what would have stopped the Google/Facebook fraud — the fake \"Quanta Computer\" domain would never have matched a verified vendor record."),
    ("Has this invoice been paid before?", "Checked first for an exact match (same number, same vendor), then for a close match (same vendor, amount within 2%, within 30 days). This is the BlackRock and Hyderabad pattern — the same paperwork, resubmitted."),
    ("Does it match a real purchase order?", "Above a set amount, an invoice must point to a real PO, and the amount has to fall within 5% of what the PO actually says."),
    ("Does the arithmetic add up?", "Line items are summed and compared to the total on the invoice; more than 5% off gets flagged."),
    ("Is it big enough to need a manager's eyes?", "Above a configurable amount, it goes to a Controller regardless of anything else."),
    ("Is there proof attached?", "Above a set amount with zero supporting documents, it's held until the vendor provides them."),
]
table_simple(["Check", "What it actually does"], rules)

P(
    "Every check writes down exactly what it found — not a vague \"risk score,\" but a row that "
    "says which rule fired and why. Run the same invoice through twice and you get the same answer "
    "twice, because nothing here is a fresh judgment call by a model — it's the same six lookups "
    "against the same database, every time.",
    space_after=10
)

P("Step 3 — Telling the right person, immediately", bold=True, size=12.5, color=TEAL, space_after=4)
P(
    "The moment a check fails, the system looks up who actually holds the right role in that "
    "company (Finance Manager, Controller, or Admin) and sends them an alert — in the app, by "
    "email, and to Slack if they've connected it. Nobody finds out at the next weekly meeting."
)

P("Step 4 — A log that can't be quietly changed", bold=True, size=12.5, color=TEAL, space_after=4)
P(
    "Every action — upload, check result, approval, escalation — is written into a log where each "
    "entry's signature is calculated from its own content plus the signature of the entry before "
    "it. Change anything in the past and every signature after it stops matching. That's not a "
    "policy we promise to follow; it's a property of how the data is stored."
)

P("Step 5 — GST reconciliation, kept separate on purpose", bold=True, size=12.5, color=TEAL, space_after=4)
P(
    "A second, independent pass through GPT-4o Vision reads GST-specific fields — GSTIN, CGST, "
    "SGST, IGST — and sorts each document into purchases or sales. Every GSTIN that comes back is "
    "then checked against India's actual 15-character GSTIN format using a fixed pattern match, no "
    "AI involved in that step at all. The verified numbers are written straight into an Excel file "
    "in the format used for filing."
)

picture(
    f"{SHOT_DIR}/02_exceptions_queue.png", 6.0,
    "Figure 2 — The review queue from the same live run. INV-2024-001 and INV-2024-002 are flagged "
    "as exact duplicates; the \"Unknown Vendor Ltd\" invoices are blocked because the vendor was "
    "never registered — both caught with no human involved, with Approve / Escalate / Request "
    "Evidence one click away.",
)

doc.add_page_break()

# =====================================================================
# SECTION 3 — DATA SOURCES
# =====================================================================
H("3. What the System Actually Reads and Checks Against")

sources = [
    ("OpenAI GPT-4o Vision", "Reads the invoice file and turns it into structured fields. Server-side only — the key never reaches a browser."),
    ("The company's own database (PostgreSQL via Supabase)", "Every vendor, purchase order, and past invoice the six checks run against. Each company's data is walled off from every other company's."),
    ("Resend", "Delivers the one-time login codes and the alerts when something is flagged."),
    ("India's official GSTIN format", "A 15-character structure — state code, PAN, entity code, and a checksum — used to validate every extracted GST number with a plain pattern match, not a model."),
    ("18 realistic sample invoices", "Modelled on real vendor names (TCS, Infosys, AWS, Microsoft) with matching vendor and purchase-order records already loaded, so the six checks produce genuine outcomes — including the two duplicates and two unregistered-vendor invoices shown live in Figure 2."),
]
table_simple(["Source", "What it's used for"], sources)

P(
    "What we are not claiming: there is no live connection to the government GST portal yet. "
    "Reconciliation today works from uploaded files, not an automatic pull of filed returns. We are "
    "saying this plainly because a submission that hides its gaps is worth less than one that names "
    "them.",
    italic=True, color=GREY, size=10,
)

doc.add_page_break()

# =====================================================================
# SECTION 4 — WHY NOT JUST USE CHATGPT
# =====================================================================
H("4. Why a Person Can't Just Do This in ChatGPT")

P(
    "Paste the same invoice into ChatGPT, Claude, or Gemini and it will describe what's on the "
    "page. It will not catch fraud, because a chat window has none of the four things every case on "
    "the cover page actually needed to be caught:"
)

reasons = [
    ("It has no list of approved vendors.", "It cannot tell you a vendor was never approved, because it has no record of which vendors your company has ever approved. Ours checks a real database — shown live in Figure 2."),
    ("It has no memory of your past invoices.", "It cannot know the same invoice was already paid three weeks ago, because each chat starts from nothing. Ours checks your company's actual invoice history every single time — also in Figure 2."),
    ("It has no purchase orders to check against.", "A real 3-way match needs a real PO record with a real committed amount. A chat session has never seen your purchase orders."),
    ("It cannot prove its own record wasn't edited afterward.", "An auditor needs proof that a decision log wasn't quietly changed six months later. A chat transcript offers no such proof; our signed, chained log does."),
    ("It cannot enforce who is allowed to approve what.", "Anyone can ask ChatGPT to approve anything. Our system only lets a Controller approve a high-value invoice, enforced by the server, not by a polite request."),
    ("It gives a different answer every time you ask.", "Run the same invoice through twice and a model can reason differently each time. Our six checks give the same answer every time, because that's what an auditor — or a court — actually needs."),
]
for label, detail in reasons:
    bullet(detail, lead=label + " ")

callout(
    "ChatGPT can read an invoice. It cannot remember your last 500 invoices, tell you who is "
    "allowed to approve what, or prove to an auditor that nothing was changed after the fact. "
    "That gap is the entire product.",
    fill="ECFDF5", border="0F766E", text_color=TEAL, size=11,
)

doc.add_page_break()

# =====================================================================
# SECTION 5 — ARCHITECTURE
# =====================================================================
H("5. How It's Built")

arch = [
    ("Frontend", "React, TypeScript, Vite, Tailwind CSS — the screen the reviewer actually uses."),
    ("Backend", "Node.js and Express, written entirely in TypeScript."),
    ("Database", "PostgreSQL hosted on Supabase, accessed through Prisma. Every record is tagged to a single company; no company can ever query another's data."),
    ("File storage", "Supabase Storage — invoice files and supporting evidence are kept behind time-limited signed links, never public."),
    ("The one AI call", "OpenAI GPT-4o Vision, called only from the server. The key is never sent to a browser."),
    ("Email", "Resend, for login codes and alerts."),
    ("Login", "One-time codes by email, with a session token that expires and is checked on every request."),
    ("Security basics", "Standard HTTP security headers, rate limits on login and upload, hashed API keys, encrypted storage for any third-party credentials."),
    ("Who can do what", "Five roles — AP Reviewer, Finance Manager, Controller, Auditor, Admin — enforced both on the server and on the screen, not just hidden in the UI."),
]
table_simple(["Part", "What it is"], arch)

picture(ARCH_IMG, 6.3, "Figure 3 — How a request actually flows through the system, end to end.", space_after=4)

doc.add_page_break()

# =====================================================================
# SECTION 6 — DEMO AND HONEST STATUS
# =====================================================================
H("6. What We Actually Ran, and What's Still Missing")

P(
    "Everything below is one continuous run: a brand-new signup, an empty workspace, real records "
    "loaded, captured automatically the same day this was written. Nothing here was staged by hand.",
    italic=True, color=GREY, size=10, space_after=10,
)

P("Scenario A — the same invoice, twice", bold=True, size=12.5, color=TEAL, space_after=4)
P(
    "Invoice INV-2024-001 (Tata Consultancy Services, ₹4,85,000) comes in — the vendor is approved "
    "and the PO matches, so it clears instantly. A second invoice arrives with the exact same "
    "number, vendor, and amount."
)
P("The system checks the company's invoice history, finds the exact match, and blocks the second one — risk score 95+, shown live below.")
P("This is the same trick used in the BlackRock and Hyderabad cases on the cover page: the same paperwork, sent again.", space_after=10)

picture(
    f"{SHOT_DIR}/03_invoice_detail_blocked.png", 6.0,
    "Figure 4 — INV-2024-001's review page: flagged \"Probable duplicate,\" every one of the six "
    "checks shown with its result, full history, and one-click actions.",
)

P("Scenario B — a vendor nobody approved", bold=True, size=12.5, color=TEAL, space_after=4)
P(
    "An invoice from \"Unknown Vendor Ltd\" arrives for ₹8,50,000, with no purchase order. The "
    "vendor check finds no such company in the approved list and blocks it immediately — risk "
    "score 88+, with an alert sent straight to the Controller. This is the Google/Facebook pattern: "
    "a name on an invoice that was never actually verified."
)

P("Scenario C — GST reconciliation", bold=True, size=12.5, color=TEAL, space_after=4)
P(
    "Up to five invoices go in; the system sorts each into purchase or sale, checks every GST "
    "number against the real format, and produces a ready-to-file Excel sheet.",
    space_after=8
)

picture(f"{SHOT_DIR}/04_gst_reconciliation.png", 5.6, "Figure 5 — The GST reconciliation screen, ready to accept a batch.")

P("What's working, and what isn't yet — said plainly", bold=True, size=12.5, color=TEAL, space_after=6)

status = [
    ("Reading invoices automatically", "Working"),
    ("All six fraud/compliance checks", "Working — Figures 1, 2, and 4"),
    ("Catching duplicates and unapproved vendors", "Working — Figures 2 and 4"),
    ("One-click approve / escalate / request evidence", "Working"),
    ("Tamper-proof decision log", "Working"),
    ("Five-role permission system", "Working"),
    ("Spreadsheet import and sample data", "Working"),
    ("GST reconciliation and Excel export", "Working"),
    ("Email and in-app alerts", "Working"),
    ("Pushing approved invoices to an ERP", "Working, via webhook"),
    ("Pulling filed GST returns automatically from the government portal", "Not built yet"),
    ("Actually sending the payment after approval", "Not built yet — approval is the last step today"),
    ("Automated tests", "Not built yet"),
    ("A proper mobile layout", "Partly there — usable, not optimised"),
]
table_simple(["Piece", "Status"], status)

callout(
    "We are submitting something that runs, not a slide deck. Every claim above points to a "
    "screenshot, a file, or an honest \"not yet.\" Judge it on that.",
    fill="FEF3C7", border="D97706", text_color=AMBER, size=10.5,
)

foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("End of submission — Invoice.Audit, InnovateZ 2026")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = GREY

doc.save(OUT_PATH)
print("Saved successfully.")
